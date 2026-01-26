"""
Content extraction utilities for various file formats
Supports: DOCX, PDF, XLSX, TXT, PPTX, and more

Author: Edgar McOchieng
"""

import os
from pathlib import Path
from typing import Dict, Optional, Any
import docx
from PyPDF2 import PdfReader
import openpyxl
from datetime import datetime
from config.logger import get_logger

logger = get_logger(__name__)


class ExtractionError(Exception):
    """Raised when content extraction fails"""
    pass


class ContentExtractor:
    """
    Extract text content and metadata from various file formats
    
    Supported formats:
    - Text: .txt, .md, .log
    - Documents: .docx, .pdf
    - Spreadsheets: .xlsx, .csv
    - Presentations: .pptx (future)
    """
    
    # Mapping of extensions to extraction methods
    EXTRACTORS = {
        '.txt': '_extract_txt',
        '.md': '_extract_txt',
        '.log': '_extract_txt',
        '.docx': '_extract_docx',
        '.pdf': '_extract_pdf',
        '.xlsx': '_extract_xlsx',
    }
    
    def __init__(self):
        """Initialize the content extractor"""
        self.stats = {
            "total_extracted": 0,
            "failed": 0,
            "by_type": {}
        }
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            ExtractionError: If extraction fails
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.EXTRACTORS:
            raise ExtractionError(f"Unsupported file type: {ext}")
        
        extractor_method = getattr(self, self.EXTRACTORS[ext])
        
        try:
            logger.debug(f"Extracting content from {file_path}")
            content = extractor_method(file_path)
            
            # Update statistics
            self.stats["total_extracted"] += 1
            self.stats["by_type"][ext] = self.stats["by_type"].get(ext, 0) + 1
            
            logger.debug(f"Extracted {len(content)} characters from {os.path.basename(file_path)}")
            return content
            
        except Exception as e:
            self.stats["failed"] += 1
            logger.error(f"Failed to extract content from {file_path}: {e}")
            raise ExtractionError(f"Error extracting content from {file_path}: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing metadata
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        # Get file system metadata
        metadata = self._get_filesystem_metadata(file_path)
        
        # Get document-specific metadata
        try:
            if ext == '.docx':
                doc_metadata = self._extract_docx_metadata(file_path)
                metadata.update(doc_metadata)
            elif ext == '.pdf':
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif ext == '.xlsx':
                xlsx_metadata = self._extract_xlsx_metadata(file_path)
                metadata.update(xlsx_metadata)
        except Exception as e:
            logger.warning(f"Could not extract document metadata from {file_path}: {e}")
        
        return metadata
    
    def _get_filesystem_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract file system metadata"""
        stat = os.stat(file_path)
        
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_extension": os.path.splitext(file_path)[1].lower(),
            "file_size_bytes": stat.st_size,
            "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "accessed_time": datetime.fromtimestamp(stat.st_atime).isoformat(),
        }
        
        # Try to get Windows file owner
        try:
            import win32security
            sd = win32security.GetFileSecurity(file_path, win32security.OWNER_SECURITY_INFORMATION)
            owner_sid = sd.GetSecurityDescriptorOwner()
            name, domain, type = win32security.LookupAccountSid(None, owner_sid)
            metadata["owner"] = f"{domain}\\{name}"
        except:
            metadata["owner"] = "Unknown"
        
        return metadata
    
    # =========================================================================
    # Text Extraction Methods
    # =========================================================================
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors ignored
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + table_text
        return '\n'.join(all_text)
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        with open(file_path, 'rb') as f:
            pdf = PdfReader(f)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
            
            return '\n\n'.join(text_parts)
    
    def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX file"""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        all_text = []
        for sheet in wb.worksheets:
            all_text.append(f"[Sheet: {sheet.title}]")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    all_text.append(" | ".join(row_text))
        
        return '\n'.join(all_text)
    
    # =========================================================================
    # Metadata Extraction Methods
    # =========================================================================
    
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        doc = docx.Document(file_path)
        props = doc.core_properties
        
        metadata = {}
        
        if props.title:
            metadata["document_title"] = props.title
        if props.author:
            metadata["document_author"] = props.author
        if props.subject:
            metadata["document_subject"] = props.subject
        if props.keywords:
            metadata["document_keywords"] = props.keywords
        if props.category:
            metadata["document_category"] = props.category
        if props.comments:
            metadata["document_comments"] = props.comments
        if props.created:
            metadata["document_created"] = props.created.isoformat()
        if props.modified:
            metadata["document_modified"] = props.modified.isoformat()
        
        # Count statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        with open(file_path, 'rb') as f:
            pdf = PdfReader(f)
            
            metadata = {
                "page_count": len(pdf.pages)
            }
            
            # Extract PDF metadata
            if pdf.metadata:
                info = pdf.metadata
                
                if info.get('/Title'):
                    metadata["document_title"] = info['/Title']
                if info.get('/Author'):
                    metadata["document_author"] = info['/Author']
                if info.get('/Subject'):
                    metadata["document_subject"] = info['/Subject']
                if info.get('/Keywords'):
                    metadata["document_keywords"] = info['/Keywords']
                if info.get('/Creator'):
                    metadata["document_creator"] = info['/Creator']
                if info.get('/Producer'):
                    metadata["document_producer"] = info['/Producer']
                if info.get('/CreationDate'):
                    metadata["document_created"] = info['/CreationDate']
                if info.get('/ModDate'):
                    metadata["document_modified"] = info['/ModDate']
            
            return metadata
    
    def _extract_xlsx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from XLSX file"""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        props = wb.properties
        
        metadata = {
            "sheet_count": len(wb.worksheets),
            "sheet_names": [sheet.title for sheet in wb.worksheets]
        }
        
        if props.title:
            metadata["document_title"] = props.title
        if props.creator:
            metadata["document_author"] = props.creator
        if props.subject:
            metadata["document_subject"] = props.subject
        if props.keywords:
            metadata["document_keywords"] = props.keywords
        if props.category:
            metadata["document_category"] = props.category
        if props.description:
            metadata["document_description"] = props.description
        if props.created:
            metadata["document_created"] = props.created.isoformat()
        if props.modified:
            metadata["document_modified"] = props.modified.isoformat()
        
        return metadata
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get extraction statistics"""
        return self.stats.copy()
    
    def reset_statistics(self):
        """Reset extraction statistics"""
        self.stats = {
            "total_extracted": 0,
            "failed": 0,
            "by_type": {}
        }